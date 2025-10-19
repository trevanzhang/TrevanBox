#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
周报导出工具
将风控合规部周报从markdown格式导出为多种格式

功能：
- 剔除front matter信息
- 转换为Word文档格式（.docx）
- 转换为PDF文档格式（.pdf）
- 导出到用户桌面
- 保持中文格式和样式
- 支持批量导出多种格式

依赖：
- python-docx（Word导出）
- markdown（Markdown解析）
- beautifulsoup4（HTML解析）
- reportlab（PDF生成）

安装方式：
uv add python-docx markdown beautifulsoup4 reportlab
"""

import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime

def detect_encoding(file_path: str) -> str:
    """检测文件编码，参考prehandler.py实现"""
    try:
        import chardet
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result.get('encoding', 'utf-8')
            # 处理常见的中文编码
            if encoding.lower() in ['gb2312', 'gbk']:
                return 'gb18030'
            return encoding or 'utf-8'
    except ImportError:
        return 'utf-8'
    except:
        return 'utf-8'

def remove_frontmatter(content: str) -> str:
    """移除YAML front matter，只保留正文"""
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            return parts[2].strip()
    return content

def get_desktop_path() -> str:
    """跨平台获取桌面路径"""
    if os.name == 'nt':  # Windows
        return os.path.join(os.path.expanduser('~'), 'Desktop')
    else:  # Linux/Mac
        return os.path.join(os.path.expanduser('~'), 'Desktop')

def generate_filename(title: str = None, format_type: str = 'docx') -> str:
    """生成导出文件名"""
    if title:
        # 从标题中提取日期
        date_match = re.search(r'(\d{4}年\d{1,2}月\d{1,2}日)', title)
        if date_match:
            date_str = date_match.group(1)
            date_str = re.sub(r'年|月|日', '-', date_str).strip('-')
            return f"风控合规部周报_{date_str}.{format_type}"

    # 使用当前日期
    current_date = datetime.now().strftime("%Y-%m-%d")
    return f"风控合规部周报_{current_date}.{format_type}"

def markdown_to_html_content(content: str) -> str:
    """使用markdown库将markdown转换为HTML"""
    try:
        import markdown
        # 启用常用扩展
        extensions = ['markdown.extensions.tables', 'markdown.extensions.fenced_code']
        html = markdown.markdown(content, extensions=extensions)
        return html
    except ImportError:
        print("错误：缺少markdown依赖包")
        print("请运行：uv add markdown")
        return ""

def parse_html_to_content_blocks(html: str) -> Tuple[str, list]:
    """解析HTML内容为结构化的内容块"""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, 'html.parser')
    content_blocks = []

    # 处理不同类型的HTML元素
    for element in soup.find_all():
        if element.name == 'h1':
            text = element.get_text(strip=True)
            content_blocks.append((text, 'Heading_1', ''))
        elif element.name == 'h2':
            text = element.get_text(strip=True)
            content_blocks.append((text, 'Heading_2', ''))
        elif element.name == 'h3':
            text = element.get_text(strip=True)
            content_blocks.append((text, 'Heading_3', ''))
        elif element.name == 'p':
            # 解析段落中的格式化内容
            text = element.get_text()
            content_blocks.append((text, 'Paragraph', element))
        elif element.name in ['ul', 'ol']:
            # 处理列表
            list_items = element.find_all('li')
            for li in list_items:
                text = li.get_text(strip=True)
                if element.name == 'ul':
                    content_blocks.append((text, 'List Bullet', ''))
                else:
                    content_blocks.append((text, 'List Number', ''))

    # 如果HTML为空或解析失败，回退到简单文本处理
    if not content_blocks:
        lines = html.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                content_blocks.append((line, 'Paragraph', ''))

    # 提取文档标题
    title = generate_filename(format_type='docx')
    if content_blocks and content_blocks[0][1] == 'Heading_1':
        title = generate_filename(content_blocks[0][0], format_type='docx')

    return title, content_blocks

def markdown_to_docx_content(content: str) -> Tuple[str, list]:
    """将markdown内容转换为docx格式"""
    # 先转换为HTML
    html = markdown_to_html_content(content)
    if not html:
        return generate_filename(format_type='docx'), []

    # 解析HTML为结构化内容块
    return parse_html_to_content_blocks(html)

def add_formatted_text(paragraph, html_element_or_text):
    """向段落添加带格式的文本"""
    try:
        from bs4 import BeautifulSoup

        if hasattr(html_element_or_text, 'find_all'):  # 这是一个BeautifulSoup元素
            # 处理HTML元素
            element = html_element_or_text

            if element.name in ['p', 'span']:
                # 处理段落或span内的格式化内容
                for child in element.children:
                    if hasattr(child, 'name'):
                        if child.name == 'strong' or child.name == 'b':
                            run = paragraph.add_run(child.get_text())
                            run.bold = True
                        elif child.name == 'em' or child.name == 'i':
                            run = paragraph.add_run(child.get_text())
                            run.italic = True
                        elif child.name == 'u':
                            run = paragraph.add_run(child.get_text())
                            run.underline = True
                        else:
                            paragraph.add_run(child.get_text())
                    else:
                        # 文本节点
                        paragraph.add_run(str(child))
            else:
                # 其他类型的HTML元素，直接添加文本
                paragraph.add_run(element.get_text())
        else:
            # 纯文本，直接添加
            paragraph.add_run(html_element_or_text)

    except ImportError:
        # 如果没有bs4，直接添加纯文本
        paragraph.add_run(str(html_element_or_text))
    except Exception as e:
        # 出错时回退到纯文本
        paragraph.add_run(str(html_element_or_text))

def create_docx(title: str, content_blocks: list, output_path: str) -> bool:
    """创建Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # 确保基本样式存在
        if 'Heading_1' not in doc.styles:
            doc.styles.add_style('Heading_1', WD_STYLE_TYPE.PARAGRAPH)
        if 'Heading_2' not in doc.styles:
            doc.styles.add_style('Heading_2', WD_STYLE_TYPE.PARAGRAPH)
        if 'Heading_3' not in doc.styles:
            doc.styles.add_style('Heading_3', WD_STYLE_TYPE.PARAGRAPH)

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)

        for text, style_name, html_element in content_blocks:
            if not text:  # 跳过空行
                continue

            if style_name == 'Heading_1':
                para = doc.add_heading(text, level=1)
            elif style_name == 'Heading_2':
                para = doc.add_heading(text, level=2)
            elif style_name == 'Heading_3':
                para = doc.add_heading(text, level=3)
            elif style_name == 'List Bullet':
                para = doc.add_paragraph(style='List Bullet')
                add_formatted_text(para, html_element or text)
            elif style_name == 'List Number':
                para = doc.add_paragraph(style='List Number')
                add_formatted_text(para, html_element or text)
            else:
                para = doc.add_paragraph()
                add_formatted_text(para, html_element or text)

            # 设置标题字体
            if style_name.startswith('Heading'):
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.font.bold = True

                if style_name == 'Heading_1':
                    for run in para.runs:
                        run.font.size = Pt(16)
                elif style_name == 'Heading_2':
                    for run in para.runs:
                        run.font.size = Pt(14)
                elif style_name == 'Heading_3':
                    for run in para.runs:
                        run.font.size = Pt(12)
            elif style_name.startswith('List'):
                # 列表项字体设置
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(11)
            else:
                # 普通段落字体设置
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(11)

        doc.save(output_path)
        return True

    except ImportError:
        print("错误：缺少python-docx依赖包")
        print("请运行：uv add python-docx")
        return False
    except Exception as e:
        print(f"创建Word文档失败：{e}")
        return False

def create_pdf(title: str, content_blocks: list, output_path: str) -> bool:
    """创建PDF文档"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        # 注册中文字体
        try:
            # 尝试注册常见的中文字体
            font_paths = [
                'C:/Windows/Fonts/simhei.ttf',  # Windows 黑体
                'C:/Windows/Fonts/simsun.ttc',  # Windows 宋体
                '/System/Library/Fonts/PingFang.ttc',  # Mac
                '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',  # Linux
            ]

            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('SimHei', font_path))
                        font_registered = True
                        break
                    except:
                        continue

            if not font_registered:
                print("警告：未找到中文字体，PDF可能无法正确显示中文")
                return False

        except Exception as e:
            print(f"字体注册失败：{e}")
            return False

        # 创建PDF文档
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()

        # 自定义样式
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Title'],
            fontName='SimHei',
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            textColor='#000000'
        )

        heading1_style = ParagraphStyle(
            'ChineseHeading1',
            parent=styles['Heading1'],
            fontName='SimHei',
            fontSize=14,
            leading=18,
            textColor='#000000'
        )

        heading2_style = ParagraphStyle(
            'ChineseHeading2',
            parent=styles['Heading2'],
            fontName='SimHei',
            fontSize=12,
            leading=16,
            textColor='#000000'
        )

        normal_style = ParagraphStyle(
            'ChineseNormal',
            parent=styles['Normal'],
            fontName='SimHei',
            fontSize=10,
            leading=14,
            textColor='#000000'
        )

        # 构建内容
        story = []

        for text, style_name, _ in content_blocks:
            if text.strip():  # 只处理非空行
                if style_name == 'Heading_1':
                    story.append(Paragraph(text, heading1_style))
                elif style_name == 'Heading_2':
                    story.append(Paragraph(text, heading2_style))
                elif style_name == 'Heading_3':
                    story.append(Paragraph(text, heading2_style))
                elif style_name.startswith('List'):
                    # 列表项，添加缩进
                    list_text = f"• {text}"
                    story.append(Paragraph(list_text, normal_style))
                else:
                    story.append(Paragraph(text, normal_style))

                story.append(Spacer(1, 6))

        # 生成PDF
        doc.build(story)
        return True

    except ImportError:
        print("错误：缺少reportlab依赖包")
        print("请运行：uv add reportlab")
        return False
    except Exception as e:
        print(f"创建PDF文档失败：{e}")
        return False

def export_to_docx(markdown_file: str, output_dir: Optional[str] = None) -> Optional[str]:
    """导出markdown文件为docx格式"""
    return export_document(markdown_file, output_dir, 'docx')

def export_to_pdf(markdown_file: str, output_dir: Optional[str] = None) -> Optional[str]:
    """导出markdown文件为pdf格式"""
    return export_document(markdown_file, output_dir, 'pdf')

def export_document(markdown_file: str, output_dir: Optional[str], format_type: str) -> Optional[str]:
    """导出markdown文件为指定格式"""

    # 检查输入文件
    if not os.path.exists(markdown_file):
        print(f"错误：文件不存在：{markdown_file}")
        return None

    # 确定输出目录
    if output_dir is None:
        output_dir = get_desktop_path()

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    try:
        # 读取文件内容
        encoding = detect_encoding(markdown_file)
        with open(markdown_file, 'r', encoding=encoding) as f:
            content = f.read()

        # 移除front matter
        clean_content = remove_frontmatter(content)

        if not clean_content.strip():
            print("错误：文件内容为空")
            return None

        # 转换内容
        title, content_blocks = markdown_to_docx_content(clean_content)

        # 生成输出文件路径和文件名
        filename = generate_filename(title if content_blocks and content_blocks[0][1] == 'Heading_1' else None, format_type)
        output_file = os.path.join(output_dir, filename)

        # 根据格式创建文档
        if format_type == 'docx':
            success = create_docx(title, content_blocks, output_file)
        elif format_type == 'pdf':
            success = create_pdf(title, content_blocks, output_file)
        else:
            print(f"错误：不支持的格式：{format_type}")
            return None

        if success:
            print(f"成功导出到：{output_file}")

            # 显示文件大小
            file_size = os.path.getsize(output_file)
            if file_size < 1024:
                size_str = f"{file_size} B"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / (1024 * 1024):.1f} MB"

            print(f"文件大小：{size_str}")
            return output_file
        else:
            return None

    except Exception as e:
        print(f"导出失败：{e}")
        return None

def export_multiple_formats(markdown_file: str, output_dir: Optional[str], formats: list) -> list:
    """导出markdown文件为多种格式"""
    results = []

    for format_type in formats:
        print(f"正在导出{format_type.upper()}格式...")
        result = export_document(markdown_file, output_dir, format_type)
        if result:
            results.append(result)
        else:
            print(f"{format_type.upper()}格式导出失败")

    return results

def main():
    parser = argparse.ArgumentParser(
        description="周报导出工具 - 将markdown周报导出为多种格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例：
  python export-to-docx.py 周报.md --format docx
  python export-to-docx.py 周报.md --format pdf
  python export-to-docx.py 周报.md --format both
  python export-to-docx.py 周报.md --output-dir ./exports
        """
    )

    parser.add_argument(
        'markdown_file',
        help='周报markdown文件路径'
    )

    parser.add_argument(
        '--format',
        choices=['docx', 'pdf', 'both'],
        default='docx',
        help='导出格式（默认：docx）'
    )

    parser.add_argument(
        '--output-dir',
        help=f'输出目录（默认为桌面：{get_desktop_path()}）'
    )

    parser.add_argument(
        '--version',
        action='version',
        version='周报导出工具 v2.0'
    )

    args = parser.parse_args()

    print("开始导出周报...")

    if args.format == 'both':
        # 导出多种格式
        results = export_multiple_formats(args.markdown_file, args.output_dir, ['docx', 'pdf'])
        if results:
            print(f"导出完成！共导出{len(results)}个文件：")
            for result in results:
                print(f"  - {result}")
            sys.exit(0)
        else:
            print("导出失败")
            sys.exit(1)
    else:
        # 导出单一格式
        result = export_document(args.markdown_file, args.output_dir, args.format)
        if result:
            print(f"导出完成！文件位置：{result}")
            sys.exit(0)
        else:
            print("导出失败")
            sys.exit(1)

if __name__ == "__main__":
    main()